#!/usr/bin/env python3
"""
清洗题库 JSON 源数据并重新生成 Word/CSV 导入模板。

用法：
    python scripts/clean_and_generate_templates.py

处理内容：
1. 天翼云：将"正确/错误"双选项单选题转换为判断题；修复 <br> 合并的选项。
2. 阿里云：手动修复 16 道选项被合并或嵌入题干的题目；其余保持原样。
3. 重新生成 docs/ 下的 .docx 与 .csv 模板文件。
"""

import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

OPTION_BR_RE = re.compile(r"<br\s*/?>\s*([A-Z])[\.．、,，\)\]\'\s]+")


def split_br_options(content: str) -> list[str] | None:
    """按 <br> + 选项标号拆分被合并的选项内容。"""
    if "<br" not in content:
        return None
    parts = OPTION_BR_RE.split(content)
    if len(parts) == 1:
        return None
    result = [parts[0].strip()]
    for i in range(1, len(parts), 2):
        if i + 1 < len(parts):
            result.append(parts[i + 1].strip())
    return result


def clean_tianyi(data: dict) -> dict:
    """清洗天翼云 JSON。"""
    for q in data["questions"]:
        opts = q.get("options", [])

        # 正确/错误 双选项 → 判断题
        if len(opts) == 2:
            contents = {o["content"].strip() for o in opts}
            if contents == {"正确", "错误"}:
                q["type"] = "judgment"

        # 修复 <br> 合并选项
        new_opts = []
        changed = False
        for opt in opts:
            split = split_br_options(opt["content"])
            if split and len(split) > 1:
                for j, part in enumerate(split):
                    new_opts.append({"label": chr(ord(opt["label"]) + j), "content": part})
                changed = True
            else:
                new_opts.append(opt)
        if changed:
            for i, opt in enumerate(new_opts):
                opt["label"] = chr(ord("A") + i)
            q["options"] = new_opts

    return data


# 阿里云 16 道需要手工修正的题目（按 JSON 数组中的 1-based 序号）
ALIYUN_MANUAL_FIXES = {
    27: {
        "stem": "阿里云的云盾安骑士可以记录远程访问服务器的源 IP 地址,并对频繁连接的可疑 IP 进行屏蔽， 日常运维过程中需要通过 功能对系统管理员常用的 IP 进行设置。",
        "options": [
            {"label": "A", "content": "安全组"},
            {"label": "B", "content": "webshell 云査杀"},
            {"label": "C", "content": "IP 白名单"},
            {"label": "D", "content": "常用登录地管理"},
        ],
    },
    60: {
        "options": [
            {"label": "A", "content": "创建 VPC 时，系统会自动为每个 VPC 创建 1 个路由器"},
            {"label": "B", "content": "删除 VPC 时，会自动删除对应的路由器"},
            {"label": "C", "content": "不支持直接创建和删除路由器"},
            {"label": "D", "content": "路由器不支持根据具体的路由条目的设置来转发网络流量"},
        ],
    },
    139: {
        "options": [
            {"label": "A", "content": "LoadBalancer"},
            {"label": "B", "content": "Listener"},
            {"label": "C", "content": "Backend Server"},
            {"label": "D", "content": "Heartbeat"},
        ],
    },
    163: {
        "stem": "阿里云的云盾数据风控可以很好地解决 WEB 应用中常见的垃圾注册、刷库等业务 风险识别的 难题，要想使用这项服务首先得进行业务数据的采集，对于 WEB 应用系统，可 以釆用 方式来 釆集信息。",
        "options": [
            {"label": "A", "content": "JavaScript"},
            {"label": "B", "content": "SDK"},
            {"label": "C", "content": "JavaScript 和 SDK"},
            {"label": "D", "content": "HTML5"},
        ],
    },
    184: {
        "options": [
            {"label": "A", "content": "停用加速域名会自己停止相关域名的解析"},
            {"label": "B", "content": "删除加速域名时，被加速域名的解析记录也会同时被自己恢复为加速前的内容"},
            {"label": "C", "content": "\"停用\"该加速域名后，该条加速域名信息仍保留，针对加速域名的请求系统不会再做自动回源处理"},
            {"label": "D", "content": "\"删除\"该加速域名后,将删除该加速域名的全部相关记录，对于仅需要暂停使用该加速域名，推荐选择\"停用\"方式"},
        ],
    },
    240: {
        "options": [
            {"label": "A", "content": "删除 VPC 时，系统会自动删除对应的路由表"},
            {"label": "B", "content": "不支持创建和删除路由表"},
            {"label": "C", "content": "每个路由器可以有多个路由表"},
            {"label": "D", "content": "路由表的路由条目会影响 VPC 中的所有云产品实例"},
        ],
    },
    307: {
        "stem": "开放系统互连参考模型(Open System Interconnect,简称 OSI)是国际标准化组织(ISO) 和国际电报电话咨询委员会(CCITT)联合制定的开放系统互连参考模型，为开放式互连信息系统提供了一种功能结构的框架。我们常用到的第七层是指 。",
        "options": [
            {"label": "A", "content": "应用层"},
            {"label": "B", "content": "表示层"},
            {"label": "C", "content": "会话层"},
            {"label": "D", "content": "传输层"},
        ],
    },
    436: {
        "options": [
            {"label": "A", "content": "木马文件检查"},
            {"label": "B", "content": "防 WEB 应用系统的密码破解"},
            {"label": "C", "content": "异地登录报警"},
            {"label": "D", "content": "防操作系统密码暴力破解"},
        ],
    },
    565: {
        "options": [
            {"label": "A", "content": "实例系列 II 的用户可以获得更大的实例规格"},
            {"label": "B", "content": "实例系列 II 相对于实例系列 I 增加了一些新的指令集，使整数和浮点运算的性能翻倍"},
            {"label": "C", "content": "ECS 实例创建以后可以进行实例系列 I 与实例系列 II 的变更"},
            {"label": "D", "content": "实例系列 II 全部为 I/O 优化实例，配合 SSD 云盘使用可以获得更高的 I/O 性能"},
        ],
    },
    745: {
        "stem": "阿里云 ECS 有两种网络类型经典网络与专用网络，这两种类型的区别分别是()。",
        "options": [
            {"label": "A", "content": "经典网络由阿里云统一规划和管理，更适合对网络易用性要求比较高的用户"},
            {"label": "B", "content": "专有网络是指用户在阿里云的基础网络内建立一个可以自定义的专有网络隔离网络"},
            {"label": "C", "content": "经典网络可以进行自定义网络,专有网络无法自定义"},
            {"label": "D", "content": "经典网络需要用户进行管理，专有网络更加的简单无需专业的网络管理"},
        ],
    },
    751: {
        "options": [
            {"label": "A", "content": "定时任务触发请求,时间为 8:00"},
            {"label": "B", "content": "报警任务触发请求,时间为 8:00"},
            {"label": "C", "content": "报警任务触发请求,时间为 8:06"},
            {"label": "D", "content": "伸缩组 7:58 被停用并再次启用后，报警任务触发请求，时间为 8:00"},
        ],
    },
    759: {
        "stem": "阿里云块存储服务提供数据块级别的存储服务，包括云盘和本地盘两种,在进行选择时需要了解二者的异同。本地盘与云盘对比来看()。",
        "options": [
            {"label": "A", "content": "云盘和本地盘均支持 SSD"},
            {"label": "B", "content": "云盘和本地盘均支持快照"},
            {"label": "C", "content": "云盘和本地盘均釆用多副本的机制保证数据可靠性"},
            {"label": "D", "content": "云盘和本地盘的性能均与其容量大小有关"},
        ],
    },
    788: {
        "options": [
            {"label": "A", "content": "实例启动自检等原因导致启动速度慢"},
            {"label": "B", "content": "忘记 SSH 密码，无法使用远程连接工具连接 ECS"},
            {"label": "C", "content": "实例误开启操作系统防火墙"},
            {"label": "D", "content": "云服务器失陷，CPU 和带宽消耗高"},
        ],
    },
    792: {
        "stem": "阿里云的资源是多地域(Region)、多可用区(Zone)部署的，通常来讲，同地域下不同可用区之间的距离大概在()范围",
        "options": [
            {"label": "A", "content": "数百米"},
            {"label": "B", "content": "数公里"},
            {"label": "C", "content": "数十公里"},
            {"label": "D", "content": "数百公里"},
        ],
    },
    881: {
        "options": [
            {"label": "A", "content": "RDS 实例支持内外网切换，不会影响其他与 RDS 实例的连接"},
            {"label": "B", "content": "在高安全模式下，RDS 支持仅内网连接、仅外网连接、内外网同时连接"},
            {"label": "C", "content": "用户可以在 RDS 管理控制台中进行连接方式的切换"},
            {"label": "D", "content": "在标准模式下，RDS 实例不支持内网、外网同时连接"},
        ],
    },
    886: {
        "options": [
            {"label": "A", "content": "无法看到文件是否被删除"},
            {"label": "B", "content": "需要提交工单"},
            {"label": "C", "content": "要定位删除的操作，无需开通任何服务只需要需要过滤 DELETE 关键字来定位"},
            {"label": "D", "content": "OSS 判断文件被访问操作等记录，需要开启 OSS 的 Logging 日志，需要访问日志中过滤 DELETE 关键字来定位"},
        ],
    },
}


def clean_aliyun(data: dict) -> dict:
    """清洗阿里云 JSON。"""
    questions = data["questions"]
    for qnum, fix in ALIYUN_MANUAL_FIXES.items():
        idx = qnum - 1
        if idx < len(questions):
            q = questions[idx]
            if "stem" in fix:
                q["stem"] = fix["stem"]
            q["options"] = fix["options"]
    return data


def json_to_word(data: dict, docx_path: Path, name: str, description: str) -> None:
    """从 JSON 生成 Word 导入模板。"""
    questions = data.get("questions", [])
    doc = Document()

    title = doc.add_heading("题库导入模板", level=1)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_paragraph(f"题库名称：{name}")
    doc.add_paragraph(f"题库描述：{description}")
    doc.add_paragraph()

    doc.add_heading("填写说明", level=2)
    doc.add_paragraph("1. 题库名称和描述填写在文档开头。")
    doc.add_paragraph("2. 使用【单选题】、【多选题】、【判断题】标记题型区域。")
    doc.add_paragraph("3. 每道题包含：题干、选项、答案、解析、难度（可选，1-3）。")
    doc.add_paragraph("4. 选项格式为：A. 内容 或 A、内容。")
    doc.add_paragraph("5. 答案格式为：答案：A 或 答案：A,B,C。")
    doc.add_paragraph("6. 判断题选项固定为 A. 正确、B. 错误。")
    doc.add_paragraph()

    type_titles = {"single": "【单选题】", "multiple": "【多选题】", "judgment": "【判断题】"}
    by_type: dict[str, list[dict]] = {"single": [], "multiple": [], "judgment": []}
    for q in questions:
        t = q.get("type", "single")
        by_type.setdefault(t, []).append(q)

    for t in ["single", "multiple", "judgment"]:
        qs = by_type.get(t, [])
        if not qs:
            continue
        doc.add_heading(type_titles[t], level=2)
        for i, q in enumerate(qs, 1):
            doc.add_paragraph(f"{i}. {q['stem']}")
            for opt in q.get("options", []):
                doc.add_paragraph(f"{opt['label']}. {opt['content']}")
            doc.add_paragraph(f"答案：{','.join(q.get('answer', []))}")
            doc.add_paragraph(f"解析：{q.get('explanation', '')}")
            doc.add_paragraph(f"难度：{q.get('difficulty', 2)}")
            doc.add_paragraph()

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def json_to_csv(data: dict, csv_path: Path) -> None:
    """从 JSON 生成 CSV 导入文件。"""
    questions = data.get("questions", [])
    max_opts = max((len(q.get("options", [])) for q in questions), default=0)

    headers = ["type", "stem", "answer", "explanation", "difficulty"]
    for i in range(max_opts):
        headers.append(f"option_{chr(ord('A') + i)}")

    with csv_path.open("w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(headers)
        for q in questions:
            row = [
                q.get("type", "single"),
                q.get("stem", ""),
                ",".join(q.get("answer", [])),
                q.get("explanation", ""),
                q.get("difficulty", 2),
            ]
            for opt in q.get("options", []):
                row.append(opt.get("content", ""))
            while len(row) < len(headers):
                row.append("")
            writer.writerow(row)


def main() -> None:
    # 天翼云
    tianyi_json_path = DOCS / "天翼云架构师_312题.json"
    tianyi_data = json.loads(tianyi_json_path.read_text(encoding="utf-8"))
    tianyi_data = clean_tianyi(tianyi_data)
    tianyi_json_path.write_text(json.dumps(tianyi_data, ensure_ascii=False, indent=2), encoding="utf-8")
    json_to_word(
        tianyi_data,
        DOCS / "天翼云高级解决方案架构师-导入模板.docx",
        name="天翼云高级解决方案架构师",
        description="Converted from 天翼云高级解决方案架构师.docx, 312 questions",
    )
    json_to_csv(tianyi_data, DOCS / "天翼云架构师_312题.csv")

    # 阿里云
    aliyun_json_path = DOCS / "aliyun_acp_899题.json"
    aliyun_data = json.loads(aliyun_json_path.read_text(encoding="utf-8"))
    aliyun_data = clean_aliyun(aliyun_data)
    aliyun_json_path.write_text(json.dumps(aliyun_data, ensure_ascii=False, indent=2), encoding="utf-8")
    json_to_word(
        aliyun_data,
        DOCS / "阿里云ACP题库云计算（899题）-导入模板.docx",
        name="阿里云ACP题库云计算",
        description="Converted from 阿里云ACP题库云计算（899题）.doc, 887 questions",
    )
    json_to_csv(aliyun_data, DOCS / "aliyun_acp_899题.csv")

    print("Templates regenerated successfully.")


if __name__ == "__main__":
    main()
