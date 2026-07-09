from pathlib import Path

from docx import Document
from pypdf import PdfReader


def extract_text(path: Path, file_type: str) -> str:
    if file_type == "pdf":
        reader = PdfReader(str(path))
        text = "\n".join(page.extract_text() or "" for page in reader.pages)
    elif file_type == "docx":
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            rows = []
            for row in table.rows:
                rows.append("\t".join(cell.text for cell in row.cells))
            parts.append("\n".join(rows))
        text = "\n\n".join(part for part in parts if part)
    else:
        text = path.read_text(encoding="utf-8")
    return text.strip()
